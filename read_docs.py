from docx import Document
import sys

def read_docx(file_path):
    try:
        doc = Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        
        # Đọc cả bảng nếu có
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
                    
        return "\n".join(full_text)
    except Exception as e:
        return f"Error: {str(e)}"

if __name__ == "__main__":
    skkn_path = r'c:\AI APP\app-diemdanh\docs\LÊ HẠNH NHÂN- TOÀN VĂN SKKN WEBAPP-2026.docx'
    mau_path = r'c:\AI APP\app-diemdanh\docs\Mẫu Phụ lục I , II (1).docx'
    
    print("--- CONTENT OF SKKN ---")
    print(read_docx(skkn_path))
    print("\n\n--- CONTENT OF MAU ---")
    print(read_docx(mau_path))
