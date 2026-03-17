from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_user_guide():
    doc = Document()

    # Style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # Title
    title = doc.add_heading('HƯỚNG DẪN SỬ DỤNG HỆ THỐNG ĐIỂM DANH', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('Tài liệu hướng dẫn chi tiết các chức năng cơ bản của webapp điểm danh dành cho Giáo viên và Cán bộ quản lý.')

    # 1. Đăng nhập
    doc.add_heading('1. Đăng nhập hệ thống bằng Google', level=1)
    doc.add_paragraph(
        'Để bắt đầu sử dụng, người dùng cần đăng nhập bằng tài khoản Google do trường cấp để đảm bảo tính bảo mật và phân quyền chính xác.'
    )
    p1 = doc.add_paragraph('Các bước thực hiện:')
    p1.add_run('\n- Bước 1: Truy cập địa chỉ web của hệ thống.')
    p1.add_run('\n- Bước 2: Bấm vào nút "Sign in with Google".')
    p1.add_run('\n- Bước 3: Chọn tài khoản email của bạn.')

    img1_path = r'C:\Users\BCD\.gemini\antigravity\brain\5e3268f5-6d7a-4c4a-bd4b-19e77e460aee\login_guide_1773638466350.png'
    if os.path.exists(img1_path):
        doc.add_picture(img1_path, width=Inches(5))
        caption = doc.add_paragraph('Ảnh 1: Giao diện đăng nhập Google')
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 2. Điểm danh
    doc.add_heading('2. Chức năng Điểm danh (5 Trường hợp)', level=1)
    doc.add_paragraph('Tại trang "Điểm danh nhanh", danh sách học sinh sẽ hiển thị kèm các nút trạng thái.')

    # Case 1
    doc.add_heading('Trường hợp 1: Học sinh có mặt (Hiện diện)', level=2)
    doc.add_paragraph('- Mô tả: Học sinh đi học đầy đủ, đúng giờ.')
    doc.add_paragraph('- Thao tác: Mặc định tất cả học sinh đều ở trạng thái có mặt. Bạn không cần thực hiện thao tác nào nếu học sinh có mặt đầy đủ.')
    
    img2_path = r'C:\Users\BCD\.gemini\antigravity\brain\5e3268f5-6d7a-4c4a-bd4b-19e77e460aee\attendance_header_guide_1773638482000.png'
    if os.path.exists(img2_path):
        doc.add_picture(img2_path, width=Inches(5))
        caption = doc.add_paragraph('Ảnh 2: Giao diện danh sách học sinh có mặt')
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Case 2
    doc.add_heading('Trường hợp 2: Vắng có phép (P)', level=2)
    doc.add_paragraph('- Mô tả: Học sinh vắng mặt và đã nộp đơn xin phép.')
    doc.add_paragraph('- Thao tác: Bấm vào nút chữ "P" (màu vàng) tương ứng với tên học sinh.')
    
    img3_path = r'C:\Users\BCD\.gemini\antigravity\brain\5e3268f5-6d7a-4c4a-bd4b-19e77e460aee\attendance_p_guide_1773638540780.png'
    if os.path.exists(img3_path):
        doc.add_picture(img3_path, width=Inches(2))
        caption = doc.add_paragraph('Ảnh 3: Trạng thái vắng có phép (P)')
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Case 3
    doc.add_heading('Trường hợp 3: Vắng không phép (K)', level=2)
    doc.add_paragraph('- Mô tả: Học sinh vắng mặt không rõ lý do hoặc không có đơn phép.')
    doc.add_paragraph('- Thao tác: Bấm vào nút chữ "K" (màu đỏ) tương ứng với tên học sinh.')
    
    img4_path = r'C:\Users\BCD\.gemini\antigravity\brain\5e3268f5-6d7a-4c4a-bd4b-19e77e460aee\attendance_k_guide_1773638560826.png'
    if os.path.exists(img4_path):
        doc.add_picture(img4_path, width=Inches(2))
        caption = doc.add_paragraph('Ảnh 4: Trạng thái vắng không phép (K)')
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Case 4
    doc.add_heading('Trường hợp 4: Đi trễ hoặc vắng tiết lẻ (T)', level=2)
    doc.add_paragraph('- Mô tả: Học sinh vào lớp muộn hoặc chỉ vắng một số tiết nhất định trong buổi học.')
    doc.add_paragraph('- Thao tác: Bấm vào nút "T" (màu xanh). Sau đó, bạn có thể chọn các số tiết tương ứng (1, 2, 3, 4, 5) mà học sinh vắng.')
    
    img5_path = r'C:\Users\BCD\.gemini\antigravity\brain\5e3268f5-6d7a-4c4a-bd4b-19e77e460aee\attendance_t_guide_1773638577769.png'
    if os.path.exists(img5_path):
        doc.add_picture(img5_path, width=Inches(4))
        caption = doc.add_paragraph('Ảnh 5: Giao diện chọn tiết vắng (T)')
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Case 5
    doc.add_heading('Trường hợp 5: Vi phạm kỷ luật (VP)', level=2)
    doc.add_paragraph('- Mô tả: Ghi nhận các trường hợp vi phạm nội quy (ví dụ: không đồng phục, mất trật tự).')
    doc.add_paragraph('- Thao tác: Bấm nút "VP" (màu tím). Bạn có thể vuốt hàng học sinh sang trái để nhập nội dung vi phạm chi tiết.')
    
    img6_path = r'C:\Users\BCD\.gemini\antigravity\brain\5e3268f5-6d7a-4c4a-bd4b-19e77e460aee\attendance_vp_guide_1773638600003.png'
    if os.path.exists(img6_path):
        doc.add_picture(img6_path, width=Inches(5))
        caption = doc.add_paragraph('Ảnh 6: Bảng nhập thông tin vi phạm')
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 3. Báo cáo
    doc.add_heading('3. Xem và Xuất báo cáo dữ liệu', level=1)
    doc.add_paragraph(
        'Hệ thống cung cấp công cụ báo cáo mạnh mẽ để theo dõi tình hình chuyên cần của học sinh theo thời gian.'
    )
    doc.add_paragraph('Các thao tác chính:')
    doc.add_paragraph('1. Chuyển sang menu "Báo cáo".', style='List Bullet')
    doc.add_paragraph('2. Chọn khoảng thời gian (Từ ngày - Đến ngày) và Lớp cần xem.', style='List Bullet')
    doc.add_paragraph('3. Kiểm tra số liệu thống kê hiển thị trên màn hình.', style='List Bullet')
    doc.add_paragraph('4. Bấm nút "Xuất file" (biểu tượng Excel) để tải dữ liệu về máy phục vụ công tác lưu trữ hoặc in ấn.', style='List Bullet')

    img7_path = r'C:\Users\BCD\.gemini\antigravity\brain\5e3268f5-6d7a-4c4a-bd4b-19e77e460aee\report_guide_1773638498706.png'
    if os.path.exists(img7_path):
        doc.add_picture(img7_path, width=Inches(5))
        caption = doc.add_paragraph('Ảnh 7: Giao diện Dashboard và Xuất báo cáo')
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Save
    output_path = r'C:\Users\BCD\.gemini\antigravity\brain\5e3268f5-6d7a-4c4a-bd4b-19e77e460aee\User_Guide_Attendance_App.docx'
    doc.save(output_path)
    print(f"File saved successfully at: {output_path}")

if __name__ == "__main__":
    create_user_guide()
