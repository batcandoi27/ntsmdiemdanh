from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_phu_luc():
    doc = Document()

    # Cấu hình font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)

    # Header Phụ lục I
    p_header = doc.add_paragraph()
    p_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_header.add_run('CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM\nĐộc lập - Tự do - Hạnh phúc')
    run.bold = True

    doc.add_paragraph('\n')

    # Phụ lục I
    h1 = doc.add_heading('Phụ lục I', level=1)
    h1.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_title.add_run('ĐƠN YÊU CẦU CÔNG NHẬN SÁNG KIẾN')
    run.bold = True
    run.font.size = Pt(16)

    doc.add_paragraph('Kính gửi: Ban Giám hiệu Trường THCS Trần Bội Cơ')

    p = doc.add_paragraph()
    p.add_run('Tôi là tác giả: ').bold = True
    p.add_run('LÊ HẠNH NHÂN')

    p = doc.add_paragraph()
    p.add_run('Đề nghị công nhận sáng kiến: ').bold = True
    p.add_run('Số hóa công tác quản lý chuyên cần và kỷ luật học sinh thông qua WebApp tại trường THCS TRẦN BỘI CƠ')

    p = doc.add_paragraph()
    p.add_run('Đã áp dụng/áp dụng thử từ ngày: ').bold = True
    p.add_run('01/09/2025')
    p.add_run(' tại: ').bold = True
    p.add_run('Trường THCS Trần Bội Cơ')

    p = doc.add_paragraph()
    p.add_run('Hiệu quả chính: ').bold = True
    p.add_run('\n- Giảm 80% thời gian tổng hợp báo cáo chuyên cần.\n- Tăng tính chính xác và minh bạch trong việc theo dõi kỷ luật.\n- Kết nối thông tin tức thời giữa giáo viên, giám thị và nhà trường.\n- Lưu trữ dữ liệu khoa học, dễ dàng tra cứu và đối soát.')

    doc.add_paragraph('\nTôi xin cam đoan mọi thông tin nêu trong đơn là trung thực, đúng sự thật và hoàn toàn chịu trách nhiệm trước pháp luật.')

    # Footer Phụ lục I
    p_footer = doc.add_paragraph()
    p_footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_footer.add_run('\nTP. Hồ Chí Minh, ngày 16 tháng 03 năm 2026\n')
    p_footer.add_run('Người yêu cầu công nhận\n\n\n')
    run = p_footer.add_run('LÊ HẠNH NHÂN')
    run.bold = True

    doc.add_page_break()

    # Phụ lục II
    h2 = doc.add_heading('Phụ lục II', level=1)
    h2.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    p_title2 = doc.add_paragraph()
    p_title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_title2.add_run('BẢN MÔ TẢ NỘI DUNG CƠ BẢN CỦA SÁNG KIẾN')
    run.bold = True
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    p.add_run('Tên Sáng kiến: ').bold = True
    p.add_run('Số hóa công tác quản lý chuyên cần và kỷ luật học sinh thông qua WebApp tại trường THCS TRẦN BỘI CƠ')

    p = doc.add_paragraph()
    p.add_run('Tác giả: ').bold = True
    p.add_run('LÊ HẠNH NHÂN')

    # Mục 1: Thực trạng
    p = doc.add_paragraph()
    p.add_run('1. Thực trạng:').bold = True
    p.add_run('\n- Công tác điểm danh và quản lý kỷ luật truyền thống bằng sổ tay tốn nhiều thời gian.\n- Dữ liệu rải rác, khó tổng hợp báo cáo nhanh các trường hợp vi phạm thường xuyên.\n- Thiếu sự đồng bộ thông tin kịp thời giữa các bộ phận quản lý.\n- Quá trình tra cứu lịch sử chuyên cần của học sinh mất nhiều công sức.')

    # Mục 2: Nội dung sáng kiến
    p = doc.add_paragraph()
    p.add_run('2. Nội dung sáng kiến:').bold = True
    p.add_run('\n- Xây dựng WebApp chuyên dụng cho công tác điểm danh nhanh trên thiết bị di động.\n- Thiết kế hệ thống cơ sở dữ liệu (Firebase/Supabase) để lưu trữ tập trung.\n- Tự động hóa quy trình xuất báo cáo thống kê định kỳ.\n- Tích hợp các tính năng ghi nhận vi phạm, khen thưởng và theo dõi tiết lẻ.')

    # Mục 3: Hiệu quả mang lại
    p = doc.add_paragraph()
    p.add_run('3. Hiệu quả mang lại:').bold = True
    p.add_run('\n- Tối ưu hóa quy trình quản lý, giúp giáo viên tập trung vào chuyên môn.\n- Nâng cao ý thức kỷ luật của học sinh nhờ sự theo dõi sát sao.\n- Cải thiện hiệu quả phối hợp trong nhà trường thông qua dữ liệu thời gian thực.\n- Tiết kiệm chi phí in ấn văn phòng phẩm truyền thống.')

    doc.add_paragraph('\nĐánh giá phạm vi ảnh hưởng của Sáng kiến: [x] Chỉ có hiệu quả trong phạm vi Đơn vị áp dụng')

    # Footer Phụ lục II
    p_footer2 = doc.add_paragraph()
    p_footer2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_footer2.add_run('\nTP. Hồ Chí Minh, ngày 16 tháng 03 năm 2026\n')
    p_footer2.add_run('Người yêu cầu công nhận\n\n\n')
    run = p_footer2.add_run('LÊ HẠNH NHÂN')
    run.bold = True

    # Lưu file
    output_path = r'c:\AI APP\app-diemdanh\docs\Phu_luc_I_II_LE_HANH_NHAN_Final.docx'
    doc.save(output_path)
    print(f"File created: {output_path}")

if __name__ == "__main__":
    create_phu_luc()
